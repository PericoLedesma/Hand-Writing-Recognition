import numpy as np
import cv2
from os import listdir
import matplotlib.pyplot as plt
import statistics

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt,RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def smooth(x,window_len=10,window='flat'):
    return np.convolve(x, np.ones(window_len)/window_len, mode='valid')

def get_hist(bin_img, type="binary_sum"):
    hist = []
    if type == "binary_sum":
        hist = np.sum(bin_img, axis=0)
        hist = smooth(hist)
    elif type == "connected_comp":
        for row in bin_img.T:
            num_labels, labels_im = cv2.connectedComponents(row)
            hist.append(num_labels)
    return hist

def extract_rois(horizontal_hist, THRESHOLD = 1):
    ROIs_idx = []
    idxs = []

    for idx, ele in enumerate(horizontal_hist):

        if ele <= THRESHOLD and idx != len(horizontal_hist) - 1:
            idxs.append(idx)
        else:
            if idxs:
                ROIs_idx.append(idxs)
                idxs = []

    return ROIs_idx

def whitespace_deletion(image):
    up_bound = 0
    low_bound = len(image)
    left_bound = 0
    right_bound = len(image[0])
    #vertical crop
    for idx in range(0, len(image[0])):
        if np.any(image[:, idx]):
            left_bound = idx
            break
    for idx in range(len(image[0]) - 1, -1, -1):
        if np.any(image[:, idx]):
            right_bound = idx
            break

    #horizontal crop
    for idx in range(0, len(image)):
        if np.any(image[idx, :]):
            up_bound = idx
            break
    for idx in range(len(image) - 1, -1, -1):
        if np.any(image[idx, :]):
            low_bound = idx
            break
    return image[up_bound:low_bound, left_bound:right_bound]

def preprocess_image(image):
    imgray = cv2.cvtColor(image, cv2.COLOR_BGR2GRAY)
    filteredImg = cv2.medianBlur(imgray, ksize=11)
    ret, thresh1 = cv2.threshold(filteredImg,127,255,cv2.THRESH_BINARY)
    return thresh1

def extract_characters(image, ROIs, line_nr, name):
    characters = []

    for idx in range(0, len(ROIs) - 1):
        segmentation_left = int(ROIs[idx][-1])
        segmentation_right = int(ROIs[idx+1][0])
        char = image[:, segmentation_left:segmentation_right]
        cropped_char = whitespace_deletion(char)
        count = np.sum(cropped_char == 255)
        if count > 0:
            characters.append(cropped_char)
            #cv2.rectangle(image, (segmentation_left, 0), (segmentation_right, len(image)), (255,0,0), 2)
    #cv2.imwrite("results/" + str(name) + "_" + str(line_nr) + "_.jpg", image)
    return characters

'''
file_name = "P632-Fg002-R-C01-R01-binarized_line_"
for line in range(0,8):
    sentence_img = cv2.imread("results/" + file_name + str(line) + ".jpg")

    thresh1 = preprocess_image(sentence_img)

    horizontal_hist = get_hist(thresh1, type="connected_comp")
    ROIs = extract_rois(horizontal_hist)

    chars = extract_characters(thresh1, ROIs)

    for idx, char in enumerate(chars):
        cv2.imwrite("chars/" + str(idx) + "_" + str(line) + ".jpg", char)
'''
def output_doc(texts, name):
    '''
    texts: the text list. e.g. [[ab],[cd]]. [ab] is the first line
    name: the name of the img without format (e.g. '.jpg').
    '''  
    document = Document()
    document.styles['Normal'].font.name = u'Habbakuk'
    document.styles['Normal'].font.size = Pt(12)
    document.styles['Normal'].font.color.rgb = RGBColor(0,0,0)
    for i in range(len(texts)):
        line=texts[i]
        if line==['#']:
            doc_line=document.add_paragraph('')
        else:
            doc_line=document.add_paragraph(line)
        doc_line.paragraph_format.alignment= WD_ALIGN_PARAGRAPH.RIGHT
    document.save(name+'.docx')
    print(name+'.docx'+" generated!!!")

def output_txt(texts, name):
    '''
    texts: the text list. e.g. [[ab],[cd]]. [ab] is the first line
    name: the name of the img without format (e.g. '.jpg').
    '''  
    with open(name+'.txt','wb') as f:
        for i in range(len(texts)):
            line=texts[i][0]
            if line=='#':
                f.write(str('\n').encode(encoding='utf-8'))
            else:
                f.write(str(line + '\n').encode(encoding='utf-8'))
    f.close()
    print(name+'.txt'+" generated!!!")
